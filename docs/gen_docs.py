"""
生成 PRD 和技术开发文档（Word 格式）
运行: .venv/bin/python docs/gen_docs.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = datetime.date.today().strftime("%Y年%m月%d日")
OUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ─── 样式工具 ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_para(doc, text, bold=False, indent=0, color=None, size=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.8 + 0.5)
    p.add_run(text)
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(level * 0.8 + 0.5)
    p.add_run(text)
    return p


def add_table(doc, headers, rows, header_bg="1F3864", header_fg="FFFFFF"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(header_fg)
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "EBF2FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()
    return table


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 灰色背景代码框（用带边框段落模拟）
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# PRD — 产品需求文档
# ═══════════════════════════════════════════════════════════════════════════════

def build_prd():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    # 封面
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("企业差旅报销审批智能体")
    tr.font.size = Pt(26)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("产品需求文档（PRD）")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"版本: V1.0    日期: {TODAY}    状态: POC 阶段").font.size = Pt(10)

    doc.add_page_break()

    # ── 1. 产品概述 ───────────────────────────────────────────────────────────
    add_heading(doc, "1. 产品概述", 1)

    add_heading(doc, "1.1 产品背景", 2)
    add_para(doc,
        "当前企业差旅报销流程依赖人工逐一审核，存在以下痛点：",
    )
    for pain in [
        "审批周期长：从提交到入账平均需要 5~10 个工作日",
        "规则执行不一致：审批人对政策理解差异导致同类申请处理结果不同",
        "员工提交前无预检：缺票、超标问题只在审核阶段才发现，退回重提效率低",
        "数据孤岛：报销数据与预算系统、OA 系统未打通，统计滞后",
        "人工审核重复劳动：约 80% 的申请为标准合规申请，可自动化处理",
    ]:
        add_bullet(doc, pain)

    add_heading(doc, "1.2 产品定位", 2)
    add_para(doc,
        "本产品是基于 AI 大语言模型（LLM）的差旅报销智能审批系统，通过三级 Agent 流水线实现报销申请的自动化预审、智能初审和财务终审，"
        "并通过对话式人机界面支持员工、经理、财务人员与 AI Agent 进行实时交互。"
        "POC 阶段聚焦核心审批流程验证，后续将与企业 OA 系统深度集成。"
    )

    add_heading(doc, "1.3 产品目标", 2)
    add_table(doc,
        ["目标维度", "当前现状", "目标值", "衡量方式"],
        [
            ["自动化率",      "0%（全人工）",   "≥80% 申请自动处理",   "AUTO_APPROVED 比例"],
            ["审批周期",      "5~10 工作日",     "合规申请 <1 分钟",    "提交到决定时间"],
            ["规则一致性",    "人工差异较大",    "100% 基于规则库决策",  "规则引用覆盖率"],
            ["员工体验",      "提交后被动等待",  "提交前 AI 预检通过",  "退回率降低 60%"],
            ["数据可追溯性",  "纸质/邮件记录",   "全流程决策日志留档",  "审计日志完整率"],
        ]
    )

    doc.add_page_break()

    # ── 2. 用户群体 ───────────────────────────────────────────────────────────
    add_heading(doc, "2. 用户角色与场景", 1)

    add_heading(doc, "2.1 用户角色定义", 2)
    add_table(doc,
        ["角色", "职责", "使用场景", "核心诉求"],
        [
            ["员工（申请人）",  "提交差旅报销申请",    "出差归来，发起报销",        "流程简单、快速获得结果"],
            ["部门经理",        "初审下属申请",        "批量查看并审批团队申请",    "快速了解申请合规性，一键决策"],
            ["财务专员",        "终审合规性与预算",    "核查票据合规、预算余额",    "规则引用准确、预算实时可查"],
            ["财务总监/HR",     "人工复核边缘案例",    "处理超预算/大额/特殊申请", "清晰的 Agent 分析报告"],
            ["系统管理员",      "维护规则库与配置",    "更新报销标准、城市等级",   "配置灵活、无需改代码"],
        ]
    )

    add_heading(doc, "2.2 核心用户旅程", 2)

    for journey in [
        ("员工报销旅程",
         ["1. 出差归来，打开报销助手（Agent 1）",
          "2. 通过对话描述出差信息（目的地、日期、费用）",
          "3. Agent 实时查询规定，提示超标项和缺票风险",
          "4. 查看右侧草稿面板确认信息，点击「提交」",
          "5. 系统自动判断：合规小额 → 自动通过；否则 → 进入审批流水线",
          "6. 通过消息通知跟踪审批状态（OA 集成后）"]),
        ("经理审批旅程",
         ["1. 收到待审批通知（OA 消息/邮件，集成后）",
          "2. 打开「经理初审助手」，查看待审申请列表",
          "3. 对特定申请提问：「帮我审批 EXP0082」",
          "4. Agent 自动：查申请详情 → 查规定 → 输出带规则引用的决定",
          "5. 经理确认或修改决定，系统执行并流转到财务队列"]),
        ("财务终审旅程",
         ["1. 在「财务审批助手」查看待终审申请",
          "2. 询问：「审批 EXP0082，查一下销售部预算」",
          "3. Agent 自动：查预算余额 → 查规定 → 综合给出带预算说明的终审决定",
          "4. 大额或预算预警申请标记人工复核",
          "5. 财务总监在「人工审核」页面做最终裁定"]),
    ]:
        add_heading(doc, journey[0], 3)
        for step in journey[1]:
            add_bullet(doc, step)

    doc.add_page_break()

    # ── 3. 功能需求 ───────────────────────────────────────────────────────────
    add_heading(doc, "3. 功能需求", 1)

    add_heading(doc, "3.1 三级智能审批流水线", 2)
    add_para(doc, "系统按以下顺序处理每条报销申请：", bold=False)
    add_table(doc,
        ["阶段", "处理方", "触发条件", "输出", "处理时间"],
        [
            ["规则引擎预审", "纯 Python 规则引擎",  "新申请提交",            "AUTO_APPROVED / 转人工流程",  "<1 秒"],
            ["Agent 1 预检", "报销材料预审助手",     "员工主动发起",          "草稿 + 合规报告",             "5~15 秒"],
            ["Agent 2 初审", "部门经理初审智能体",   "PENDING_MANAGER 状态",  "APPROVED / REJECTED / 人工",  "15~30 秒"],
            ["Agent 3 终审", "财务审批智能体",        "PENDING_FINANCE 状态",  "APPROVED / REJECTED / 人工",  "15~30 秒"],
            ["人工复核",     "财务总监/HR",           "PENDING_HUMAN_REVIEW",  "最终 APPROVED / REJECTED",    "人工处理"],
        ]
    )

    add_heading(doc, "3.2 自动审批规则（规则引擎）", 2)
    add_para(doc, "满足以下全部条件时，申请自动通过，无需 LLM 介入：")
    add_table(doc,
        ["条件", "阈值", "依据规则"],
        [
            ["申请总金额",      "< 1,000 元",                        "rule_006"],
            ["所有费用有发票",  "100% 有发票",                       "rule_004"],
            ["住宿费合规",      "一线≤600元/晚，二线≤400，其他≤300", "rule_002"],
            ["餐饮费合规",      "≤150元/天",                         "rule_003"],
            ["提交时限",        "出差结束后30天内",                  "rule_011"],
        ]
    )

    add_heading(doc, "3.3 Agent 1 — 报销发起助手", 2)
    add_table(doc,
        ["功能模块", "描述", "优先级"],
        [
            ["政策问答",      "员工输入任意报销政策问题，Agent 调用 RAG 查询规定库后回答",              "P0"],
            ["对话式填单",    "引导员工逐步提供申请人、目的地、日期、费用等信息",                      "P0"],
            ["实时合规预检",  "收集费用信息后，自动对照城市标准提示超标项",                            "P0"],
            ["草稿实时渲染",  "右侧面板实时展示申请草稿（申请人/费用明细/合计/问题列表）",             "P0"],
            ["一键提交",      "草稿确认后一键提交进入审批流水线，生成申请编号",                        "P0"],
            ["历史案例参考",  "遇到特殊情况时检索相似历史申请案例，供员工参考",                        "P1"],
            ["快捷问题按钮",  "预设常用问题（住宿标准/发起申请/发票处理）",                           "P1"],
        ]
    )

    add_heading(doc, "3.4 Agent 2 — 部门经理初审智能体", 2)
    add_table(doc,
        ["功能模块", "描述", "优先级"],
        [
            ["待审申请列表",    "实时展示 PENDING_MANAGER 队列，含合规预警标记",                      "P0"],
            ["对话式审批",      "经理输入「审批 EXP0082」，Agent 自动查详情→查规定→给决定",          "P0"],
            ["规则引用决策",    "决定理由中必须引用具体规则编号（rule_001~rule_012）",                "P0"],
            ["政策问答",        "回答经理关于报销政策的任何问题",                                     "P0"],
            ["历史案例参考",    "检索相似申请的历史处理结果",                                         "P1"],
            ["批量审批指令",    "经理说「批量审核所有待审申请」，Agent 逐条处理",                     "P1"],
            ["决定记录面板",    "右侧实时展示本次已做出的决定（绿/红/紫色卡片）",                    "P1"],
        ]
    )

    add_heading(doc, "3.5 Agent 3 — 财务审批智能体", 2)
    add_table(doc,
        ["功能模块", "描述", "优先级"],
        [
            ["待终审申请列表",  "实时展示 PENDING_FINANCE 队列，含经理初审结果",                     "P0"],
            ["预算实时查询",    "调用预算工具查询部门当月/年度预算余额及使用率",                      "P0"],
            ["对话式终审",      "财务输入审批指令，Agent 自动：查预算→查规定→给终审决定",            "P0"],
            ["预算预警",        "使用率>85% 时标记 PENDING_HUMAN_REVIEW",                            "P0"],
            ["合规性终审",      "对票据合规、金额合理性做最终核查",                                   "P0"],
            ["政策问答",        "回答财务人员关于合规、预算管理的问题",                               "P1"],
        ]
    )

    add_heading(doc, "3.6 人工审核", 2)
    add_table(doc,
        ["功能模块", "描述", "优先级"],
        [
            ["待审列表",    "展示所有 PENDING_HUMAN_REVIEW 申请，含 Agent 标记理由",   "P0"],
            ["申请详情",    "查看完整费用明细、审批历史、Agent 决定理由",               "P0"],
            ["一键批准/拒绝", "填写审核意见后一键操作，结果实时写入系统",             "P0"],
            ["审核意见必填", "不填写意见无法完成操作，确保审计可追溯",                "P0"],
        ]
    )

    add_heading(doc, "3.7 RAG 知识库检索", 2)
    add_table(doc,
        ["知识库", "数据来源", "更新方式", "用途"],
        [
            ["规则库 (expense_rules)",       "data/rules/rule_001~012.txt",     "管理员修改 txt 文件后重建向量库",  "Agent 查规定、政策问答"],
            ["历史申请库 (expense_applications)", "data/applications/EXP*.txt", "每次申请决定后自动追加",          "相似案例参考"],
        ]
    )

    doc.add_page_break()

    # ── 4. 申请状态流转 ───────────────────────────────────────────────────────
    add_heading(doc, "4. 申请状态流转", 1)

    add_para(doc, "系统定义 7 种申请状态，流转规则如下：")
    add_table(doc,
        ["状态", "说明", "下一步状态", "操作方"],
        [
            ["PENDING_AUTO_CHECK",   "新申请待规则引擎处理",   "AUTO_APPROVED 或 PENDING_MANAGER",     "规则引擎"],
            ["AUTO_APPROVED",        "规则引擎自动通过",        "（终态）",                             "规则引擎"],
            ["PENDING_MANAGER",      "待部门经理初审",          "PENDING_FINANCE / REJECTED / 人工",    "Agent 2 / 经理"],
            ["PENDING_FINANCE",      "待财务终审",              "APPROVED / REJECTED / 人工",            "Agent 3 / 财务"],
            ["PENDING_HUMAN_REVIEW", "待人工复核",              "APPROVED / REJECTED",                  "财务总监 / HR"],
            ["APPROVED",             "最终批准",                "（终态）",                             "Agent 3 / 人工"],
            ["REJECTED",             "拒绝退回",                "（终态，员工可修改后重新提交）",        "任意 Agent / 人工"],
        ]
    )

    doc.add_page_break()

    # ── 5. OA 集成规划 ────────────────────────────────────────────────────────
    add_heading(doc, "5. OA 系统集成规划（POC 后）", 1)

    add_heading(doc, "5.1 集成方式", 2)
    add_table(doc,
        ["集成点", "集成方式", "说明"],
        [
            ["申请数据同步",   "REST API / Webhook",  "OA 新申请提交时触发 Webhook，写入本系统"],
            ["审批结果回写",   "REST API",            "Agent 决定后调用 OA API 更新申请状态"],
            ["消息通知",       "企微/钉钉 Webhook",   "审批完成后向申请人、经理发送消息通知"],
            ["单点登录",       "SSO / OAuth 2.0",    "用户身份与 OA 系统统一"],
            ["预算数据同步",   "定时同步 / API 拉取", "从财务系统获取实时预算数据"],
            ["文件附件",       "对象存储 / CDN",      "发票图片/PDF 存储与审核"],
        ]
    )

    add_heading(doc, "5.2 集成阶段规划", 2)
    add_table(doc,
        ["阶段", "里程碑", "预期周期"],
        [
            ["Phase 0 (当前)", "POC 验证：三级 Agent + Streamlit UI 人机交互",     "已完成"],
            ["Phase 1",        "OA 申请数据接入：Webhook 触发，自动进入审批流水线", "2~3 周"],
            ["Phase 2",        "审批结果回写 OA + 消息通知（企微/钉钉）",           "2~3 周"],
            ["Phase 3",        "SSO 登录 + 角色权限管理",                           "1~2 周"],
            ["Phase 4",        "实时预算系统对接 + 规则管理后台",                   "3~4 周"],
            ["Phase 5",        "发票 OCR 识别 + 自动核验金额",                      "4~6 周"],
        ]
    )

    doc.add_page_break()

    # ── 6. 非功能需求 ─────────────────────────────────────────────────────────
    add_heading(doc, "6. 非功能需求", 1)
    add_table(doc,
        ["类别", "指标", "目标值"],
        [
            ["性能",    "规则引擎响应时间",     "< 100ms"],
            ["性能",    "Agent 响应时间",        "< 30 秒（含 LLM + RAG）"],
            ["可用性",  "系统可用性",            "≥ 99.5%（工作时间）"],
            ["安全",    "API Key 管理",          "不允许硬编码在代码中（通过环境变量配置）"],
            ["安全",    "数据隐私",              "员工申请数据不出企业私有环境"],
            ["可扩展性","规则库更新",            "管理员修改 txt 文件即可生效，无需重新部署"],
            ["可追溯",  "审批日志",              "每步操作记录时间戳、操作人、理由"],
            ["兼容性",  "浏览器",                "Chrome / Edge / Safari 最新版"],
        ]
    )

    doc.add_page_break()

    # ── 7. 成功标准 ───────────────────────────────────────────────────────────
    add_heading(doc, "7. POC 成功标准", 1)
    add_table(doc,
        ["验收项目", "验收标准"],
        [
            ["三级 Agent 流水线",    "100 条 Mock 申请按预期比例通过各阶段，结果与设计一致"],
            ["对话式交互",           "三个 Agent 均可通过自然语言对话完成审批/问答"],
            ["规则引用准确性",       "Agent 决定理由中正确引用规则编号，无幻觉数字"],
            ["RAG 检索准确性",       "政策问答结果与 rule_*.txt 原文一致"],
            ["人工审核流程",         "PENDING_HUMAN_REVIEW 申请可通过 UI 批准/拒绝"],
            ["数据可视化",           "仪表板正确展示实时统计、漏斗图、金额分布"],
            ["OA 集成接口就绪",      "提供清晰的 REST API 接口文档供 OA 系统对接"],
        ]
    )

    # 页脚
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run(f"本文档由企业报销审批智能体 POC 团队生成  |  {TODAY}  |  版本 V1.0")
    footer_r.font.size = Pt(9)
    footer_r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    path = os.path.join(OUT_DIR, "PRD_企业报销审批智能体.docx")
    doc.save(path)
    print(f"✓ PRD 已生成: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# 技术开发文档
# ═══════════════════════════════════════════════════════════════════════════════

def build_tech_doc():
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    # 封面
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("企业差旅报销审批智能体")
    tr.font.size = Pt(26)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("技术开发文档")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"版本: V1.0    日期: {TODAY}    适用人员: 后端 / AI / 全栈工程师").font.size = Pt(10)

    doc.add_page_break()

    # ── 1. 技术栈总览 ─────────────────────────────────────────────────────────
    add_heading(doc, "1. 技术栈总览", 1)

    add_table(doc,
        ["层次", "技术/框架", "版本", "用途"],
        [
            ["LLM 推理",      "千问 qwen-plus (阿里百炼)",  "API",         "Agent 对话推理、function calling"],
            ["Embedding",     "text-embedding-v3 (DashScope)", "1024-dim", "文本向量化，用于 Milvus 检索"],
            ["Agent 框架",    "CrewAI",                     "≥0.80",       "批处理 Agent（经理/财务审批）"],
            ["对话引擎",      "OpenAI SDK (兼容接口)",      "≥1.0",        "Chat Agent 1/2/3 多轮对话"],
            ["向量数据库",    "Milvus Lite (pymilvus)",     "≥2.4",        "规则库和申请库 RAG 检索"],
            ["UI 框架",       "Streamlit",                  "≥1.30",       "人机交互 POC Web 界面"],
            ["数据可视化",    "Plotly",                     "≥5.0",        "漏斗图、饼图、直方图"],
            ["数据处理",      "Pandas",                     "≥2.0",        "表格展示"],
            ["Word 生成",     "python-docx",                "≥1.1",        "文档生成"],
            ["任务调度",      "APScheduler",                "≥3.10",       "定时批处理审批（生产环境）"],
            ["Python",        "CPython",                    "3.13",        "主运行环境"],
        ]
    )

    doc.add_page_break()

    # ── 2. 项目结构 ───────────────────────────────────────────────────────────
    add_heading(doc, "2. 项目目录结构", 1)

    add_code_block(doc, """\
expense_reimbursement/
├── demo.py                    # 命令行批量运行入口（含 --no-llm 模式）
├── ui.py                      # Streamlit Web UI（8 个页面）
├── pipeline.py                # 审批流水线（规则引擎 + Agent 调用）
├── requirements.txt           # Python 依赖
├── .env.example               # 环境变量模板
│
├── agents/
│   ├── __init__.py
│   └── definitions.py         # 三个 CrewAI Agent 定义（prep/manager/finance）
│
├── crews/
│   ├── __init__.py
│   ├── prep_crew.py           # Agent 1 批处理模式（CrewAI）
│   ├── manager_crew.py        # Agent 2 批处理模式（CrewAI）
│   ├── finance_crew.py        # Agent 3 批处理模式（CrewAI）
│   └── chat_crew.py           # 三个 Agent 对话模式（OpenAI SDK + function calling）
│
├── config/
│   └── company_rules.py       # 城市等级配置、金额阈值常量
│
├── data/
│   ├── mock_applications.py   # Mock 数据生成器（读取 txt 文件）
│   ├── rules/
│   │   ├── rule_001.txt       # 差旅日补贴规定
│   │   ├── rule_002.txt       # 住宿费规定
│   │   └── ... (共12条)
│   └── applications/
│       ├── EXP0001.txt        # Mock 申请（结构化 txt）
│       └── ... (共100条)
│
├── rag/
│   ├── expense_rag.py         # 规则库 RAG（expense_rules 集合）
│   ├── application_rag.py     # 申请库 RAG（expense_applications 集合）
│   └── expense_rules.db       # Milvus Lite 本地数据库（gitignore）
│
├── state/
│   ├── store.py               # ApplicationStore（内存 + JSON 持久化）
│   └── applications.json      # 运行时状态（gitignore）
│
├── tools/
│   ├── policy_tool.py         # CrewAI Tool：RAG 规则检索
│   └── budget_tool.py         # CrewAI Tool：部门预算查询
│
├── scheduler/
│   └── cron_jobs.py           # APScheduler 定时任务（生产环境）
│
└── docs/
    ├── gen_docs.py            # 本文档生成脚本
    ├── PRD_企业报销审批智能体.docx
    └── 技术开发文档_企业报销审批智能体.docx""")

    doc.add_page_break()

    # ── 3. 核心模块说明 ───────────────────────────────────────────────────────
    add_heading(doc, "3. 核心模块详细说明", 1)

    add_heading(doc, "3.1 状态管理 (state/store.py)", 2)
    add_para(doc, "ApplicationStore 是系统的核心状态容器，管理所有申请的生命周期。")
    add_table(doc,
        ["方法", "参数", "说明"],
        [
            ["load_applications(apps)",           "list[dict]",          "批量导入申请，初始状态 PENDING_AUTO_CHECK"],
            ["get(app_id)",                        "str",                 "按编号获取申请记录"],
            ["get_by_status(status)",              "str",                 "按状态批量查询"],
            ["update_status(app_id, new_status, actor, reason)", "...",  "状态变更，自动追加历史记录"],
            ["approve_auto(app_id, reason)",       "str, str",           "规则引擎自动批准"],
            ["manager_approve/reject/flag_human",  "str, str",           "经理初审三种操作"],
            ["finance_approve/reject/flag_human",  "str, str",           "财务终审三种操作"],
            ["save(path)",                         "str",                 "持久化到 JSON 文件"],
            ["stats()",                            "—",                   "返回各状态计数 dict"],
        ]
    )
    add_para(doc, "注：POC 阶段使用内存字典存储，生产环境应替换为 PostgreSQL / MySQL。", color="666666")

    add_heading(doc, "3.2 审批流水线 (pipeline.py)", 2)
    add_para(doc, "三个核心函数对应三个审批阶段：")
    add_code_block(doc, """\
# 步骤 1: 规则引擎（无 LLM，毫秒级）
run_auto_approve_batch(store: ApplicationStore) -> dict
  # 遍历 PENDING_AUTO_CHECK 申请，检查5个条件，决定 AUTO_APPROVED 或 PENDING_MANAGER

# 步骤 2: 经理初审（CrewAI 批处理模式）
process_manager_review(app: dict, store: ApplicationStore) -> str
  # 调用 manager_crew.run_manager_review(app)，解析 APPROVED/REJECTED/PENDING_HUMAN_REVIEW

# 步骤 3: 财务终审（CrewAI 批处理模式）
process_finance_review(app: dict, store: ApplicationStore, manager_decision: dict) -> str
  # 调用 finance_crew.run_finance_review(app, manager_decision)，同上解析""")

    add_heading(doc, "3.3 对话 Agent 引擎 (crews/chat_crew.py)", 2)
    add_para(doc, "基于 OpenAI SDK + function calling 的多轮对话引擎，三个 Agent 共用底层架构：")
    add_table(doc,
        ["函数", "角色", "可用工具", "返回值"],
        [
            ["chat(history, user_message)",
             "Agent 1 报销发起助手",
             "search_policy / search_similar_cases / check_budget",
             "(reply: str, draft: dict | None)"],
            ["chat_manager(history, user_message, pending_apps, apps_map)",
             "Agent 2 经理初审智能体",
             "search_policy / get_app_detail / search_similar_cases / make_decision",
             "(reply: str, decisions: list[dict])"],
            ["chat_finance(history, user_message, pending_apps, apps_map)",
             "Agent 3 财务审批智能体",
             "search_policy / check_budget / get_app_detail / search_similar_cases / make_decision",
             "(reply: str, decisions: list[dict])"],
        ]
    )

    add_para(doc, "工具调用循环（最多 5 次）：")
    add_code_block(doc, """\
for _ in range(5):
    response = client.chat.completions.create(model="qwen-plus", tools=TOOLS, ...)
    if msg.tool_calls:
        # 执行工具（RAG/预算/详情查询），结果追加到 messages
        continue
    # 无工具调用 → 返回最终回复
    return reply, decisions""")

    add_heading(doc, "3.4 RAG 模块", 2)

    add_heading(doc, "3.4.1 规则库 (rag/expense_rag.py)", 3)
    add_table(doc,
        ["接口", "说明"],
        [
            ["initialize(force=False)",           "从 data/rules/*.txt 读取并向量化，写入 expense_rules 集合"],
            ["search(query, top_k=3)",            "语义检索规定，返回格式化文本（含规则编号和相关度）"],
            ["search_with_ids(query, top_k=3)",   "返回结构化列表 [{rule_id, category, content, score}]"],
            ["get_rule_by_id(rule_id)",           "精确按 rule_id 读取对应 txt 文件，返回规定全文"],
            ["get_all_rules_summary()",           "返回全部规定一览（供 Agent system prompt 使用）"],
        ]
    )

    add_heading(doc, "3.4.2 申请库 (rag/application_rag.py)", 3)
    add_table(doc,
        ["接口", "说明"],
        [
            ["initialize(force=False)",           "从 data/applications/*.txt 读取并向量化（批量10条），写入 expense_applications 集合"],
            ["search_similar(query, top_k=5)",    "语义检索相似历史申请，返回格式化文本"],
            ["upsert_decision(app_id, decision, reason)", "审批完成后将决定追加写回 Milvus"],
        ]
    )

    add_heading(doc, "3.4.3 Milvus 集合配置", 3)
    add_table(doc,
        ["参数", "expense_rules", "expense_applications"],
        [
            ["维度",          "1024 (text-embedding-v3)", "1024 (text-embedding-v3)"],
            ["度量类型",      "COSINE",                   "COSINE"],
            ["动态字段",      "开启",                     "开启"],
            ["自增 ID",       "开启",                     "开启"],
            ["主要字段",      "rule_id, category, content", "app_id, applicant, department, total, outcome, content"],
            ["数据量",        "12 条",                    "100+ 条（持续增加）"],
        ]
    )

    doc.add_page_break()

    # ── 4. 数据模型 ───────────────────────────────────────────────────────────
    add_heading(doc, "4. 数据模型", 1)

    add_heading(doc, "4.1 申请记录（ApplicationStore 内部格式）", 2)
    add_code_block(doc, """\
{
  "app_id":                    "EXP0081",        # 申请编号（CHT* 为对话提交）
  "applicant":                 "张伟",
  "department":                "销售部",
  "level":                     "P5",
  "destination":               "上海",
  "trip_start":                "2026-02-28",
  "trip_end":                  "2026-03-02",
  "trip_days":                 3,
  "purpose":                   "大客户拜访",
  "expense_items": [
    {"category": "住宿费", "amount": 1500.0, "has_receipt": true,  "note": "500元/晚×3晚"},
    {"category": "餐饮费", "amount": 420.0,  "has_receipt": true,  "note": "140元/天×3天"},
    {"category": "交通费", "amount": 218.0,  "has_receipt": true,  "note": "高铁往返"}
  ],
  "total_amount":              2138.0,
  "has_all_receipts":          true,
  "all_items_compliant":       true,
  "within_limits":             true,
  "justification":             "",
  "submitted_days_after_trip": 5,
  "status":                    "APPROVED",       # 当前状态
  "history": [                                   # 审批历史
    {"timestamp": "2026-03-08T10:00:00", "from": "PENDING_AUTO_CHECK",
     "to": "PENDING_MANAGER", "actor": "auto_system", "reason": "金额超1000元"},
    {"timestamp": "2026-03-08T10:05:00", "from": "PENDING_MANAGER",
     "to": "PENDING_FINANCE", "actor": "manager_agent",
     "reason": "各项费用符合rule_002/rule_003标准，票据完整，予以初审通过"}
  ],
  "manager_decision": {"decision": "APPROVED", "reason": "..."},
  "finance_decision":  {"decision": "APPROVED", "reason": "..."}
}""")

    add_heading(doc, "4.2 规则文件格式 (data/rules/rule_XXX.txt)", 2)
    add_code_block(doc, """\
【rule_002 住宿费】
住宿费报销上限：一线城市600元/晚，二线城市400元/晚，其他城市300元/晚。
超出标准需提供说明并由部门总监审批。必须提供住宿发票。
城市分级：北京、上海、广州、深圳为一线；杭州、成都、武汉、西安等为二线。""")

    add_heading(doc, "4.3 申请文件格式 (data/applications/EXP*.txt)", 2)
    add_code_block(doc, """\
申请编号: EXP0081
申请人: 张伟
部门: 销售部
级别: P5
目的地: 上海
出差开始: 2026-02-28
出差结束: 2026-03-02
出差天数: 3
出差目的: 大客户拜访
提交天数: 5

费用明细:
  [住宿费] 1500元  有发票  500元/晚×3晚，符合一线城市标准
  [餐饮费] 420元   有发票  140元/天×3天，低于150元上限
  [交通费] 218元   有发票  高铁往返经济舱

合计: 2138元
票据完整: 是
各项合规: 是
预期结果: APPROVED""")

    doc.add_page_break()

    # ── 5. API 接口（OA 集成） ────────────────────────────────────────────────
    add_heading(doc, "5. OA 集成 API 接口设计", 1)
    add_para(doc, "以下接口供 OA 系统在 Phase 1 集成时对接，均为 RESTful JSON API：")

    add_heading(doc, "5.1 申请提交接口", 2)
    add_code_block(doc, """\
POST /api/v1/applications
Content-Type: application/json
Authorization: Bearer <token>

Request Body:
{
  "app_id":      "OA_20260308_001",   # OA 系统原始编号
  "applicant":   "张伟",
  "department":  "销售部",
  "destination": "上海",
  "trip_start":  "2026-02-28",
  "trip_end":    "2026-03-02",
  "expense_items": [...],
  "total_amount": 2138.0,
  "justification": ""
}

Response 200:
{
  "internal_id": "EXP0101",
  "status":      "PENDING_AUTO_CHECK",
  "auto_result": "PENDING_MANAGER",   # 规则引擎立即处理结果
  "message":     "申请已提交，进入经理审批队列"
}""")

    add_heading(doc, "5.2 审批状态查询接口", 2)
    add_code_block(doc, """\
GET /api/v1/applications/{app_id}/status

Response 200:
{
  "app_id":   "EXP0101",
  "status":   "APPROVED",
  "history":  [...],
  "decision": {"agent": "finance_agent", "reason": "...", "rule_refs": ["rule_002"]}
}""")

    add_heading(doc, "5.3 Webhook 回调（审批完成通知）", 2)
    add_code_block(doc, """\
POST <OA_CALLBACK_URL>   # OA 系统提供的回调地址
Content-Type: application/json

Payload:
{
  "event":      "approval_completed",
  "app_id":     "OA_20260308_001",
  "status":     "APPROVED",
  "decided_by": "finance_agent",
  "reason":     "费用合规，预算充足（剩余85%），依据rule_002/rule_003批准",
  "timestamp":  "2026-03-08T10:30:00Z"
}""")

    add_heading(doc, "5.4 规则库更新接口", 2)
    add_code_block(doc, """\
POST /api/v1/admin/rules/rebuild
Authorization: Bearer <admin_token>

Response 200:
{
  "message": "规则库已重建",
  "rule_count": 12,
  "collection": "expense_rules"
}""")

    doc.add_page_break()

    # ── 6. 部署指南 ───────────────────────────────────────────────────────────
    add_heading(doc, "6. 部署指南", 1)

    add_heading(doc, "6.1 环境准备", 2)
    add_table(doc,
        ["依赖", "最低要求", "推荐"],
        [
            ["Python",       "3.11+",      "3.13"],
            ["内存",         "4 GB",       "8 GB（含 Milvus Lite）"],
            ["磁盘",         "2 GB",       "10 GB（含向量数据库）"],
            ["网络",         "可访问阿里云 DashScope API", "固定 IP（API 白名单）"],
            ["操作系统",     "Linux / macOS", "Ubuntu 22.04 / macOS 14+"],
        ]
    )

    add_heading(doc, "6.2 安装步骤", 2)
    add_code_block(doc, """\
# 1. 克隆代码
git clone https://github.com/muguangde/expense_reimbursement_agent.git
cd expense_reimbursement_agent

# 2. 创建虚拟环境
python3.13 -m venv .venv
source .venv/bin/activate          # Linux/macOS
# .venv\\Scripts\\activate          # Windows

# 3. 安装依赖
pip install -r requirements.txt
pip install streamlit plotly python-docx

# 4. 配置 API Key（复制 .env.example 并填写）
cp .env.example .env
# 编辑 .env，填入 QWEN_API_KEY

# 5. 初始化 Milvus 向量数据库（首次运行必须）
python -c "
from rag.expense_rag import initialize as init_rules
from rag.application_rag import initialize as init_apps
init_rules(force=True)
init_apps(force=True)
print('RAG 初始化完成')
"

# 6. 运行 Web UI
streamlit run ui.py --server.port 8501

# 7. 或运行命令行批量演示
python demo.py --no-rag-init""")

    add_heading(doc, "6.3 环境变量配置", 2)
    add_code_block(doc, """\
# .env
QWEN_API_KEY=sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx

# 可选：覆盖默认 URL
# DASHSCOPE_BASE_URL=https://dashscope.aliyuncs.com/compatible-mode/v1""")
    add_para(doc, "注意：生产环境中，API Key 必须通过环境变量或 Secrets Manager 管理，禁止硬编码在源代码中。", color="CC0000")

    add_heading(doc, "6.4 生产环境建议", 2)
    add_table(doc,
        ["组件", "POC 方案", "生产替换方案"],
        [
            ["状态存储",  "内存 + JSON 文件",  "PostgreSQL / MySQL + Redis（缓存）"],
            ["向量数据库","Milvus Lite (本地)", "Milvus 集群 / 阿里云向量检索服务"],
            ["Web 服务",  "Streamlit",          "FastAPI + Vue/React 前端"],
            ["任务调度",  "手动触发",           "APScheduler / Celery + RabbitMQ"],
            ["认证授权",  "无",                 "OAuth 2.0 / SSO（与 OA 系统集成）"],
            ["日志监控",  "控制台输出",         "ELK Stack / 阿里云日志服务"],
            ["API Key",   "硬编码（POC）",      "阿里云 KMS / Vault"],
        ]
    )

    doc.add_page_break()

    # ── 7. 关键设计决策 ───────────────────────────────────────────────────────
    add_heading(doc, "7. 关键设计决策与注意事项", 1)

    add_heading(doc, "7.1 LLM 幻觉防控", 2)
    add_para(doc, "问题：LLM 可能臆造报销标准（如「经理级住宿450元/晚」），与实际规定不符。")
    add_para(doc, "解决方案：", bold=True)
    for item in [
        "Task 描述中预计算并显式注入精确数字（hotel_night * days），禁止 Agent 使用其他数字",
        "约束语言：「住宿上限只能用 X 元，禁止使用任何其他数字」",
        "RAG 强制先查询：Agent 工具调用顺序设计为「先查规定，再做决定」",
        "Mock 数据严格生成：确保测试申请金额在合规范围内，避免测试数据本身误导 LLM",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.2 对话 vs CrewAI 双模式", 2)
    add_table(doc,
        ["模式", "技术实现", "适用场景", "优缺点"],
        [
            ["对话模式", "OpenAI SDK + function calling + 对话历史",
             "Agent 1/2/3 的 UI 交互页面",
             "优：多轮对话、实时响应；缺：单次处理，无批量优化"],
            ["批处理模式", "CrewAI Crew（Task 链，sequential process）",
             "demo.py 批量审批、scheduler 定时任务",
             "优：结构化任务流、易于批量；缺：无对话上下文"],
        ]
    )

    add_heading(doc, "7.3 数据文件化设计", 2)
    add_para(doc,
        "所有规则和 Mock 申请均存储为 txt 文件（data/rules/*.txt, data/applications/*.txt），而非硬编码在 Python 中。"
        "好处：")
    for b in [
        "规则更新无需改代码，只需修改 txt 文件后重建向量库",
        "申请数据与代码分离，便于独立管理和审计",
        "txt 文件可被 Milvus 读取向量化，也可被 Python 解析为结构化数据",
        "便于 Git 追踪规则变更历史",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "7.4 Milvus 双集合设计", 2)
    add_table(doc,
        ["集合", "用途", "更新频率"],
        [
            ["expense_rules",        "12 条公司规定 → Agent 查规定、政策问答",       "规则变更时（低频）"],
            ["expense_applications", "100+ 条历史申请 → 相似案例参考、上下文增强", "每次审批完成后（高频）"],
        ]
    )

    doc.add_page_break()

    # ── 8. 常见问题 ───────────────────────────────────────────────────────────
    add_heading(doc, "8. 常见问题与解决方案", 1)
    add_table(doc,
        ["问题", "原因", "解决方案"],
        [
            ["ModuleNotFoundError: crewai",
             "crewai 安装在 .venv，使用了系统 Python",
             "使用 .venv/bin/python 或 .venv/bin/streamlit run"],
            ["Milvus 集合不存在",
             "首次运行未初始化",
             "执行 initialize(force=True) 或添加 --no-rag-init 参数跳过"],
            ["LLM 返回错误金额",
             "LLM 幻觉（规则理解偏差）",
             "检查 Task 描述中是否包含显式数字约束；检查 RAG 是否正常查询"],
            ["Streamlit 页面不刷新",
             "session_state 中旧数据残留",
             "点击「重置演示数据」按钮，或重启 Streamlit"],
            ["API 调用超时",
             "网络延迟或 DashScope 限流",
             "检查网络；适当降低并发；检查 API Key 配额"],
            ["向量检索结果不准确",
             "Embedding 质量或检索参数问题",
             "增大 top_k；检查 text-embedding-v3 是否正常调用"],
        ]
    )

    # 页脚
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run(f"技术开发文档  |  企业报销审批智能体 V1.0  |  {TODAY}")
    footer_r.font.size = Pt(9)
    footer_r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    path = os.path.join(OUT_DIR, "技术开发文档_企业报销审批智能体.docx")
    doc.save(path)
    print(f"✓ 技术开发文档已生成: {path}")
    return path


# ─── 主程序 ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    print("正在生成文档...")
    p1 = build_prd()
    p2 = build_tech_doc()
    print(f"\n✅ 两份文档已生成:")
    print(f"   {p1}")
    print(f"   {p2}")
